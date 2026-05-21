import pytest
import docx
from unittest.mock import patch, MagicMock

from src.l2_interfaces.host.os.skills.files.documents import HostOSDocuments
from src.l2_interfaces.host.os.client import HostOSAccessLevel


@pytest.mark.asyncio
async def test_read_document_unsupported_ext(os_client):
    doc_skill = HostOSDocuments(os_client)

    test_file = os_client.sandbox_dir / "secret.txt"
    test_file.touch()

    res = await doc_skill.read_document("sandbox/secret.txt")

    assert res.is_success is False
    assert "is not supported by this skill" in res.message


@pytest.mark.asyncio
async def test_read_document_path_traversal(os_client):
    os_client.access_level = HostOSAccessLevel.SANDBOX
    doc_skill = HostOSDocuments(os_client)

    res = await doc_skill.read_document("../../../etc/passwords.pdf")

    assert res.is_success is False
    assert "Access is permitted strictly inside sandbox/" in res.message


@pytest.mark.asyncio
async def test_read_document_docx_success(os_client):
    doc_skill = HostOSDocuments(os_client)
    test_file = os_client.sandbox_dir / "report.docx"

    doc = docx.Document()
    doc.add_paragraph("JAWL Document Reader.")
    doc.save(str(test_file))

    res = await doc_skill.read_document("sandbox/report.docx")

    assert res.is_success is True
    assert "JAWL Document Reader." in res.message


@pytest.mark.asyncio
@patch("src.l2_interfaces.host.os.skills.files.documents.pypdf")
async def test_read_document_pdf_pagination(mock_pypdf, os_client):
    doc_skill = HostOSDocuments(os_client)
    test_file = os_client.sandbox_dir / "book.pdf"
    test_file.touch()

    mock_reader = MagicMock()
    mock_page1 = MagicMock()
    mock_page1.extract_text.return_value = "Page 1."
    mock_page2 = MagicMock()
    mock_page2.extract_text.return_value = "Page 2."
    mock_page3 = MagicMock()
    mock_page3.extract_text.return_value = "Page 3."

    mock_reader.pages = [mock_page1, mock_page2, mock_page3]
    mock_pypdf.PdfReader.return_value = mock_reader

    res_full = await doc_skill.read_document("sandbox/book.pdf")
    assert res_full.is_success is True
    assert "Page 1" in res_full.message
    assert "Page 3" in res_full.message

    res_paginated = await doc_skill.read_document("sandbox/book.pdf", page_start=2, page_end=2)
    assert res_paginated.is_success is True
    assert "Page 1" not in res_paginated.message
    assert "Page 2" in res_paginated.message
    assert "Page 3" not in res_paginated.message

    res_oob = await doc_skill.read_document("sandbox/book.pdf", page_start=5)
    assert res_oob.is_success is False
    assert "Invalid page range" in res_oob.message


@pytest.mark.asyncio
async def test_read_document_truncation_limit(os_client):
    doc_skill = HostOSDocuments(os_client)

    os_client.config.file_read_max_chars = 50
    test_file = os_client.sandbox_dir / "long.docx"

    doc = docx.Document()
    doc.add_paragraph("A" * 500)
    doc.save(str(test_file))

    res = await doc_skill.read_document("sandbox/long.docx")

    assert res.is_success is True
    assert "Reached limit" in res.message
    assert len(res.message) < 150
